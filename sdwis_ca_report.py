# sdwis_ca_report.py
import sys
import re
import requests
import pandas as pd
import functools


# Optional: silence HTTPS warnings if your env uses unverified TLS
import urllib3
urllib3.disable_warnings(urllib3.exceptions.InsecureRequestWarning)

try:
    from docx import Document
    HAVE_DOCX = True
except Exception:
    HAVE_DOCX = False

# ----------------------- Config -----------------------

# Fields to keep per table (UPPERCASE after fetch)
TABLE_FIELDS = {
    "WATER_SYSTEM": [
        "PWSID", "PWS_ACTIVITY_CODE", "PWS_TYPE_CODE", "PWS_NAME",
        "POPULATION_SERVED_COUNT", "PRIMARY_SOURCE_CODE", "OWNER_TYPE_CODE",
        "GW_SW_CODE", "IS_GRANT_ELIGIBLE_IND", "IS_WHOLESALER_IND",
        "SERVICE_CONNECTIONS_COUNT", "ORG_NAME", "ADMIN_NAME", "EMAIL_ADDR",
        "STATE_CODE", "CITY_NAME"
    ],
    "GEOGRAPHIC_AREA": [
        "PWSID", "TRIBAL_CODE", "STATE_SERVED", "CITY_SERVED", "COUNTY_SERVED"
    ],
    "VIOLATION": [
        "PWSID", "CONTAMINANT_CODE", "VIOLATION_CODE", "VIOLATION_CATEGORY_CODE",
        "IS_HEALTH_BASED_IND", "COMPLIANCE_STATUS_CODE", "RULE_GROUP_CODE"
    ],
    "WATER_SYSTEM_FACILITY": [
        "PWSID", "FACILITY_ID", "FACILITY_NAME", "STATE_FACILITY_ID",
        "FACILITY_ACTIVITY_CODE", "FACILITY_TYPE_CODE", "IS_SOURCE_IND",
        "WATER_TYPE_CODE", "AVAILABILITY_CODE"
    ],
    "SERVICE_AREA": [
        "PWSID", "SELLER_TREATMENT_CODE", "SELLER_PWSID", "SELLER_PWS_NAME",
        "IS_SOURCE_TREATED_IND", "SERVICE_AREA_TYPE_CODE", "IS_PRIMARY_SERVICE_AREA_CODE"
    ],
    "TREATMENT": [
        "COMMENTS_TEXT", "FACILITY_ID", "PWSID", "TREATMENT_ID",
        "TREATMENT_OBJECTIVE_CODE", "TREATMENT_PROCESS_CODE"
    ]
}

# Descriptions for common codes (expand as needed)
CODE_DESCRIPTIONS = {
    # --- Summary codes ---
    "PWS_TYPE_CODE": {
        # Your requested spellings
        "CWS": "Community Water System",
        "NP": "Non-Public",
        "NTNCWS": "Non-Transient, Non-Community Water System",
        "TNCWS": "Transient, Non-Community Water System",
        "U": "Unknown",
        # Also accept legacy short forms (seen in some SDWIS pulls)
        "C": "Community Water System",
        "NTNC": "Non-Transient, Non-Community Water System",
        "NC": "Transient, Non-Community Water System",
    },
    "OWNER_TYPE_CODE": {
        "F": "Federal government",
        "L": "Local government",
        "N": "Native American",
        "P": "Private",
        "M": "Public/Private",
        "S": "State government",
    },
    "PWS_ACTIVITY_CODE": {"A": "Active", "I": "Inactive"},
    "PRIMARY_SOURCE_CODE": {
    "GW": "Ground Water",
    "SW": "Surface Water",
    "GU": "Ground water under the influence of surface water",
},
    # --- Facility codes ---
    "FACILITY_TYPE_CODE": {
        "CC": "Consecutive Connection",
        "IG": "Infiltration Gallery",
        "IN": "Intake",
        "NP": "Non-piped",
        "RC": "Roof Catchment",
        "RS": "Reservoir",
        "CW": "Clear Well",
        "ST": "Storage",
        "TP": "Treatment Plant",
        "SP": "Spring",
        "WL": "Well (Source)",
    },
    "WATER_TYPE_CODE": {
        "GW": "Ground Water",
        "SW": "Surface Water",
        "GU": "Ground water under the influence of surface water",
    },
    "AVAILABILITY_CODE": {
        "E": "Emergency",
        "I": "Interim",
        "O": "Other",
        "P": "Permanent",
        "S": "Seasonal",
        "U": "Unknown",
    },

    # --- Treatment codes ---
    "TREATMENT_OBJECTIVE_CODE": {
        "B": "DISINFECTION BY-PRODUCTS CONTROL",
        "C": "CORROSION CONTROL",
        "D": "DISINFECTION",
        "E": "DE-CHLORINATION",
        "F": "IRON REMOVAL",
        "I": "INORGANICS REMOVAL",
        "M": "MANGANESE REMOVAL",
        "O": "ORGANICS REMOVAL",
        "P": "PARTICULATE REMOVAL",
        "R": "RADIONUCLIDES REMOVAL",
        "S": "SOFTENING (HARDNESS REMOVAL)",
        "T": "TASTE / ODOR CONTROL",
        "Z": "OTHER",
    },

    # --- Violations ---
    "VIOLATION_CATEGORY_CODE": {
        "MCL": "Maximum Contaminant Level Violations",
        "MRDL": "Maximum Residual Disinfectant Level",
        "TT": "Treatment Technique Violations",
        "MR": "Monitoring and Reporting Violations",
    },
    "VIOLATION_CODE": {
        # Current
        "01": "MCL, Single Sample",
        "02": "MCL, Average",
        "03": "Monitoring, Regular",
        "04": "Monitoring, Check/Repeat/Confirmation",
        "05": "Notification, State",
        "07": "Treatment Techniques",
        "08": "Variance/Exemption/Other Compliance",
        "09": "Record Keeping",
        "11": "Non-Acute MRDL",
        "12": "Qualified Operator Failure",
        "13": "Acute MRDL",
        "19": "Monitoring, GWR Assessment Source Water",
        "20": "Failure to Consult/Respond",
        "27": "Monitoring, Routine (DBP)",
        "28": "Sanitary Survey Cooperation Failure",
        "29": "Failure Submit Filter Profile/CPE Report",
        "30": "Monitoring/Submit Plan (IDSE)",
        "31": "Monitoring Treatment (SWTR-Unfilt/GWR)",
        "32": "Reporting and Source Monitoring LT2",
        "33": "Failure Submit Treatment Requirement Rpt",
        "34": "Monitoring, GWR Triggered/Additional",
        "35": "Failure Submit Stage 2 DBPR Report",
        "36": "Monitoring Treatment (Surface Filter)",
        "37": "Treatment Tech. No Prior State Approval",
        "38": "M&R Filter Turbidity Reporting",
        "39": "Monitoring and Reporting (FBR)",
        "40": "Treatment Technique (FBR)",
        "41": "Failure to Maintain Microbial Treatment",
        "42": "Failure to Provide Treatment",
        "43": "Single Turbidity Exceed (Enhanced SWTR)",
        "44": "Treatment Technique Exceeds Turb 0.3 NTU",
        "45": "Failure Address a Deficiency",
        "46": "Treatment Technique Precursor Removal",
        "47": "Treatment Technique Uncovered Reservoir",
        "48": "Failure To Address Contamination",
        "51": "Initial LCR Tap Sampling",
        "52": "Follow-up and Routine Tap Sampling",
        "53": "Water Quality Parameter M & R",
        "56": "Initial, Follow-up, or Routine SOWT M&R",
        "57": "OCCT/SOWT Recommendation",
        "58": "OCCT/SOWT Installation",
        "59": "Water Quality Parameter Non-Compliance",
        "63": "MPL Non-Compliance",
        "64": "Lead Service Line Replacement (LSLR)",
        "65": "Public Education",
        "66": "Lead Consumer Notice",
        "71": "CCR Complete Failure to Report",
        "72": "CCR Inadequate Reporting",
        "73": "Failure to Notify Other PWS",
        "75": "PN Violation for an NPDWR Violation",
        "76": "PN Violation without NPDWR Violation",
        "77": "Tier 1 PN for Lead ALE",
        "1A": "MCL, E. coli (RTCR)",
        "2A": "TT, Level 1 Assessment (RTCR)",
        "2B": "TT, Level 2 Assessment (RTCR)",
        "2C": "TT, Corrective/Expedited Actions (RTCR)",
        "2D": "Treatment Technique Startup Procedures",
        "2E": "LSL Inventory",
        "3A": "Monitoring, Routine (RTCR)",
        "3B": "Monitoring, Additional Routine (RTCR)",
        "3C": "Monitor Coliform Turbidity Trigger (RTCR)",
        "3D": "Monitoring, Lab Cert/Method Error (RTCR)",
        "4A": "Reporting, Assessment Forms (RTCR)",
        "4B": "Report Sample Result/Fail Monitor (RTCR)",
        "4C": "Report Startup Procedures Cert Form (RTCR)",
        "4D": "Notification, E Coli Positive (RTCR)",
        "4E": "Notification, E. coli MCL (RTCR)",
        "4F": "Notify L1/L2 TT Vio, Correct Action RTCR",
        "4G": "LSL Reporting",
        "4H": "LSL Notification",
        "5A": "Sample Siting Plan Errors (RTCR)",
        "5B": "Recordkeeping Violations (RTCR)",
        # Historical
        "06": "Notification, Public",
        "10": "Operations Report",
        "21": "MCL, Acute (TCR)",
        "22": "MCL, Monthly (TCR)",
        "23": "Monitoring, Routine Major (TCR)",
        "24": "Monitoring, Routine Minor (TCR)",
        "25": "Monitoring, Repeat Major (TCR)",
        "26": "Monitoring, Repeat Minor (TCR)",
    },
}


# Paging for local pulls (bump if needed)
PAGE_SIZE = 50000
MAX_PAGES = 10   # up to ~500k rows per table

BASE = "https://data.epa.gov/efservice"
PWSID_URL = BASE + "/{table}/PWSID/{pwsid}/JSON"
ROWS_URL  = BASE + "/{table}/Rows/{start}:{end}/JSON"
FILTER_URL = BASE + "/{table}/{col}/{val}/Rows/{start}:{end}/JSON"

# ----------------------- HTTP helpers -----------------------

# (Optional) persistent session w/ small retry pool
_session = requests.Session()
_adapter = requests.adapters.HTTPAdapter(pool_connections=10, pool_maxsize=10, max_retries=3)
_session.mount("https://", _adapter)
_session.mount("http://", _adapter)

def api_get_json(url: str):
    """GET with verify True then fallback to False (common in corp envs)."""
    try:
        r = _session.get(url, timeout=60)
        r.raise_for_status()
        return r.json()
    except Exception:
        r = _session.get(url, timeout=60, verify=False)
        r.raise_for_status()
        return r.json()

def df_upper(df: pd.DataFrame) -> pd.DataFrame:
    if df is None or df.empty:
        return pd.DataFrame()
    out = df.copy()
    out.columns = [c.upper() for c in out.columns]
    return out

def pull_rows_paged(table: str, page_size: int = PAGE_SIZE, max_pages: int = MAX_PAGES) -> pd.DataFrame:
    parts = []
    for p in range(max_pages):
        start = p * page_size
        end = start + page_size
        url = ROWS_URL.format(table=table, start=start, end=end)
        try:
            data = api_get_json(url)
            if not isinstance(data, list) or not data:
                break
            parts.append(pd.DataFrame(data))
        except Exception as e:
            print(f"[Rows {table} {start}:{end}] {e}")
            break
    if not parts:
        return pd.DataFrame()
    return df_upper(pd.concat(parts, ignore_index=True))

def pull_rows_filtered(table: str, col: str, val: str,
                       page_size: int = PAGE_SIZE, max_pages: int = MAX_PAGES) -> pd.DataFrame:
    """Pull rows from EPA API with a server-side filter to reduce data size."""
    parts = []
    for p in range(max_pages):
        start = p * page_size
        end = start + page_size
        url = FILTER_URL.format(table=table, col=col, val=val, start=start, end=end)
        try:
            data = api_get_json(url)
            if not isinstance(data, list) or not data:
                break
            parts.append(pd.DataFrame(data))
        except Exception as e:
            print(f"[Rows {table} {start}:{end}] {e}")
            break
    if not parts:
        return pd.DataFrame()
    return df_upper(pd.concat(parts, ignore_index=True))

def looks_like_pwsid(s: str) -> bool:
    return bool(re.fullmatch(r"[A-Za-z]{2}\d{7}", s.strip()))

def token_and_contains(series: pd.Series, tokens: list[str]) -> pd.Series:
    """AND all tokens; case-insensitive; safe on NaN."""
    mask = pd.Series(True, index=series.index)
    for t in tokens:
        if t:
            mask &= series.astype(str).str.contains(re.escape(t), case=False, na=False)
    return mask

# ------- Lightweight caches for CLI --------

@functools.lru_cache(maxsize=64)
def _ws_by_state_cached(state_code: str) -> pd.DataFrame:
    """WATER_SYSTEM filtered by state, UPPER-cased columns."""
    ws = pull_rows_filtered("WATER_SYSTEM", "STATE_CODE", state_code)
    return df_upper(ws)

@functools.lru_cache(maxsize=1)
def _ga_all_cached() -> pd.DataFrame:
    """Bulk GEOGRAPHIC_AREA once; keep common columns; UPPER-cased."""
    ga = pull_rows_paged("GEOGRAPHIC_AREA")
    ga = df_upper(ga)
    keep = [c for c in ["PWSID", "CITY_SERVED", "COUNTY_SERVED", "STATE_SERVED"] if c in ga.columns]
    ga = ga[keep] if keep else ga
    subset = [c for c in ["PWSID", "CITY_SERVED", "COUNTY_SERVED"] if c in ga.columns]
    return ga.drop_duplicates(subset=subset) if subset else ga.drop_duplicates()


# ----------------------- Search (state-aware) -----------------------

def search_by_name(state_code: str, name_query: str, county_filter: str | None) -> pd.DataFrame:
    """
    Fast CLI search:
      - WATER_SYSTEM filtered by STATE_CODE (server-side)  [cached]
      - GEOGRAPHIC_AREA pulled once and filtered locally   [cached]
      - Optional name token filter (AND across tokens)
      - Optional county/city filter against GA.COUNTY_SERVED or GA.CITY_SERVED
      - Returns a compact table with PWSID, PWS_NAME, CITY (prefers WS.CITY_NAME, else GA.CITY_SERVED), COUNTY_SERVED
    """
    sc = (state_code or "").strip().upper()
    if not re.fullmatch(r"[A-Z]{2}", sc):
        return pd.DataFrame()

    # 1) Base WATER_SYSTEM (by state)
    ws = _ws_by_state_cached(sc)
    if ws.empty or "PWSID" not in ws.columns or "PWS_NAME" not in ws.columns:
        return pd.DataFrame()

    # Keep only what we need (include CITY_NAME if present)
    keep_ws = [c for c in ["PWSID", "PWS_NAME", "CITY_NAME"] if c in ws.columns]
    ws = ws[keep_ws].drop_duplicates("PWSID") if keep_ws else ws

    # 2) Optional name filter (AND across tokens)
    q = (name_query or "").strip()
    if q and "PWS_NAME" in ws.columns:
        tokens = re.findall(r"[A-Za-z0-9]+", q)
        if tokens:
            m = token_and_contains(ws["PWS_NAME"], tokens)
            ws = ws[m]
            if ws.empty:
                return pd.DataFrame()

    # 3) If no county/city filter → return early (fast)
    if not county_filter:
        out = ws.copy()
        # Build CITY from CITY_NAME if present
        if "CITY_NAME" in out.columns:
            out["CITY"] = out["CITY_NAME"].fillna("").astype(str).str.strip()
        show = [c for c in ["PWSID", "PWS_NAME", "CITY"] if c in out.columns]
        if not show:
            show = [c for c in ["PWSID", "PWS_NAME"] if c in out.columns]
        return out[show].drop_duplicates("PWSID").sort_values("PWS_NAME").reset_index(drop=True)

    # 4) County/city filter → do a single bulk GA filter & local join
    ga = _ga_all_cached()
    if ga.empty or "PWSID" not in ga.columns:
        # No GA available, fall back to WS-only
        out = ws.copy()
        if "CITY_NAME" in out.columns:
            out["CITY"] = out["CITY_NAME"].fillna("").astype(str).str.strip()
        show = [c for c in ["PWSID", "PWS_NAME", "CITY"] if c in out.columns]
        if not show:
            show = [c for c in ["PWSID", "PWS_NAME"] if c in out.columns]
        return out[show].drop_duplicates("PWSID").sort_values("PWS_NAME").reset_index(drop=True)

    # Filter GA to the state via PWSID prefix (robust even if STATE_SERVED missing)
    ga_state = ga[ga["PWSID"].astype(str).str.startswith(sc)]

    term = county_filter.strip().lower()
    # Use Series(False, index=ga_state.index) when the column is missing
    m_county = (
        ga_state["COUNTY_SERVED"].astype(str).str.lower().str.contains(term, na=False)
        if "COUNTY_SERVED" in ga_state.columns
        else pd.Series(False, index=ga_state.index)
    )
    m_citysv = (
        ga_state["CITY_SERVED"].astype(str).str.lower().str.contains(term, na=False)
        if "CITY_SERVED" in ga_state.columns
        else pd.Series(False, index=ga_state.index)
    )
    ga_match = ga_state[m_county | m_citysv]

    if ga_match.empty:
        # Nothing matched in GA; fall back to WS-only
        out = ws.copy()
        if "CITY_NAME" in out.columns:
            out["CITY"] = out["CITY_NAME"].fillna("").astype(str).str.strip()
        show = [c for c in ["PWSID", "PWS_NAME", "CITY"] if c in out.columns]
        if not show:
            show = [c for c in ["PWSID", "PWS_NAME"] if c in out.columns]
        return out[show].drop_duplicates("PWSID").sort_values("PWS_NAME").reset_index(drop=True)

    # Join WS with GA matches to keep only systems in the county/city selection
    merge_cols = ["PWSID", "CITY_SERVED", "COUNTY_SERVED"]
    merge_cols = [c for c in merge_cols if c in ga_match.columns]
    out = ws.merge(ga_match[merge_cols], on="PWSID", how="inner")

    # Build unified CITY column: prefer WS.CITY_NAME, else GA.CITY_SERVED
    if "CITY" not in out.columns:
        out["CITY"] = ""
    if "CITY_NAME" in out.columns:
        out["CITY"] = out["CITY"].mask(out["CITY"].eq(""), out["CITY_NAME"].fillna("").astype(str).str.strip())
    if "CITY_SERVED" in out.columns:
        out["CITY"] = out["CITY"].mask(out["CITY"].eq(""), out["CITY_SERVED"].fillna("").astype(str).str.strip())

    # Final view
    show = [c for c in ["PWSID", "PWS_NAME", "CITY", "COUNTY_SERVED"] if c in out.columns]
    if not show:
        show = [c for c in ["PWSID", "PWS_NAME"] if c in out.columns]
    return out[show].drop_duplicates("PWSID").sort_values("PWS_NAME").reset_index(drop=True)

# ----------------------- Fetch tables & report -----------------------

def fetch_table_by_pwsid(table: str, pwsid: str) -> pd.DataFrame:
    url = PWSID_URL.format(table=table, pwsid=pwsid)
    try:
        data = api_get_json(url)
        df = pd.DataFrame(data) if isinstance(data, list) and data else pd.DataFrame()
        return df_upper(df)
    except Exception as e:
        print(f"[Fetch {table}] {e}")
        return pd.DataFrame()

def fetch_all_selected(pwsid: str) -> dict[str, pd.DataFrame]:
    out = {}
    for table, wanted in TABLE_FIELDS.items():
        print(f"Fetching: {table}")
        df = fetch_table_by_pwsid(table, pwsid)
        if not df.empty:
            keep = [c for c in wanted if c in df.columns]
            if keep:
                df = df[keep]
        out[table] = df
    return out

def add_code_descriptions(df: pd.DataFrame) -> pd.DataFrame:
    if df.empty:
        return df
    out = df.copy()
    for col, mapping in CODE_DESCRIPTIONS.items():
        if col in out.columns:
            out[col + "_DESC"] = out[col].map(mapping).fillna("")
    return out

def generate_report(pwsid: str, data: dict[str, pd.DataFrame], out_path: str | None = None):
    if not HAVE_DOCX:
        print("python-docx is not installed. Install it with: pip install python-docx")
        sys.exit(1)

    # ---------------- helpers ----------------
    def u(df: pd.DataFrame) -> pd.DataFrame:
        return df_upper(df)

    def get1(df: pd.DataFrame, col: str, default: str = "N/A") -> str:
        if df.empty or col not in df.columns:
            return default
        v = str(df.iloc[0].get(col, "")).strip()
        return v if v else default

    def desc(col: str, val: str | None) -> str:
        """Show 'CODE — Description' when mapping exists; otherwise raw value."""
        if val is None or (isinstance(val, float) and pd.isna(val)):
            return "N/A"
        s = str(val).strip()
        m = CODE_DESCRIPTIONS.get(col, {})
        d = m.get(s, "")
        return f"{s} — {d}" if d else s

    def yn_from(code: str | None) -> str:
        s = ("" if code is None else str(code).strip().upper())
        return {"Y": "Yes", "N": "No"}.get(s, s or "N/A")

    def active_from(code: str | None) -> str:
        s = ("" if code is None else str(code).strip().upper())
        return {"A": "Yes", "I": "No"}.get(s, s or "N/A")

    def add_table(doc, headers: list[str], rows: list[list[str]]):
        t = doc.add_table(rows=1, cols=len(headers))
        try:
            t.style = "Table Grid"
        except Exception:
            pass
        for j, h in enumerate(headers):
            t.cell(0, j).text = str(h)
        for r in rows:
            cells = t.add_row().cells
            for j, v in enumerate(r):
                cells[j].text = "" if v is None or (isinstance(v, float) and pd.isna(v)) else str(v)

    # ---------------- source tables ----------------
    ws  = u(data.get("WATER_SYSTEM", pd.DataFrame()))
    ga  = u(data.get("GEOGRAPHIC_AREA", pd.DataFrame()))
    wsf = u(data.get("WATER_SYSTEM_FACILITY", pd.DataFrame()))
    vio = u(data.get("VIOLATION", pd.DataFrame()))
    trt = u(data.get("TREATMENT", pd.DataFrame()))  # optional; used to enrich treatment rows

    # ---------------- summary ----------------
    ws_name     = get1(ws, "PWS_NAME")
    pws_type    = desc("PWS_TYPE_CODE", get1(ws, "PWS_TYPE_CODE"))
    pws_act     = desc("PWS_ACTIVITY_CODE", get1(ws, "PWS_ACTIVITY_CODE"))
    owner       = desc("OWNER_TYPE_CODE", get1(ws, "OWNER_TYPE_CODE"))
    state       = get1(ws, "STATE_CODE", pwsid[:2] if isinstance(pwsid, str) else "N/A")
    admin       = get1(ws, "ADMIN_NAME")
    email       = get1(ws, "EMAIL_ADDR")
    pop         = get1(ws, "POPULATION_SERVED_COUNT")
    svc_conn    = get1(ws, "SERVICE_CONNECTIONS_COUNT")
    primary_src = desc("PRIMARY_SOURCE_CODE", get1(ws, "PRIMARY_SOURCE_CODE"))
    wholesaler  = yn_from(get1(ws, "IS_WHOLESALER_IND"))

    # County Served: first non-empty in GA
    county = "N/A"
    if not ga.empty and "COUNTY_SERVED" in ga.columns:
        non_empty = ga["COUNTY_SERVED"].dropna().astype(str).str.strip()
        if not non_empty.empty:
            county = non_empty.iloc[0] or "N/A"

    doc = Document()
    doc.add_heading(f"Summary Information for Water Utility {pwsid}", level=0)
    doc.add_paragraph("USEPA Safe Drinking Water Information System (SDWIS)")
    doc.add_paragraph(f"Water System Name: {ws_name}")
    doc.add_paragraph(f"System Type: {pws_type}    Activity Status: {pws_act}    Ownership: {owner}")
    doc.add_paragraph(f"State: {state}    County Served: {county}")
    doc.add_paragraph(f"Administrative Contact: {admin}    Email address: {email}")
    doc.add_paragraph(f"Population Served: {pop}    Service Connections: {svc_conn}")
    doc.add_paragraph(f"Primary Source: {primary_src}    Wholesale Supplier to Other PWS’s: {wholesaler}")

    # ======================== Facilities ========================
    doc.add_heading("Facilities", level=1)

    # -------- Sources: only IS_SOURCE_IND == 'Y'; sort by type, activity, name
    doc.add_paragraph("Sources")
    src_rows = []
    if not wsf.empty and "IS_SOURCE_IND" in wsf.columns:
        src_df = wsf.copy()
        for c in ("IS_SOURCE_IND","FACILITY_TYPE_CODE","FACILITY_ACTIVITY_CODE","FACILITY_NAME"):
            if c in src_df.columns:
                src_df[c] = src_df[c].astype(str)
        src_df = src_df[src_df["IS_SOURCE_IND"].str.upper() == "Y"]
        if not src_df.empty:
            sort_cols = [c for c in ["FACILITY_TYPE_CODE","FACILITY_ACTIVITY_CODE","FACILITY_NAME"] if c in src_df.columns]
            if sort_cols:
                src_df = src_df.sort_values(sort_cols, kind="mergesort")
            for _, r in src_df.iterrows():
                src_rows.append([
                    desc("FACILITY_TYPE_CODE", r.get("FACILITY_TYPE_CODE", "")),
                    active_from(r.get("FACILITY_ACTIVITY_CODE", "")),
                    r.get("FACILITY_NAME", ""),
                    r.get("FACILITY_ID", ""),
                    r.get("STATE_FACILITY_ID", ""),
                    desc("WATER_TYPE_CODE", r.get("WATER_TYPE_CODE", "")),
                    desc("AVAILABILITY_CODE", r.get("AVAILABILITY_CODE", "")),
                ])
    if src_rows:
        add_table(doc,
                  headers=["Type", "Active?", "Name", "SDWIS Facility ID", "State Facility ID", "Water Type", "Availability"],
                  rows=src_rows)
    else:
        doc.add_paragraph("No data available.")

    # -------- Treatment: only FACILITY_TYPE_CODE == 'TP'; sort by facility_name
    doc.add_paragraph("")  # spacer
    doc.add_paragraph("Treatment")
    tr_rows = []
    if not wsf.empty and "FACILITY_TYPE_CODE" in wsf.columns:
        tp_df = wsf.copy()
        for c in ("FACILITY_TYPE_CODE","FACILITY_NAME","FACILITY_ACTIVITY_CODE"):
            if c in tp_df.columns:
                tp_df[c] = tp_df[c].astype(str)
        tp_df = tp_df[tp_df["FACILITY_TYPE_CODE"].str.upper() == "TP"]
        if not tp_df.empty:
            # enrich with TREATMENT (objective/process) if available
            if not trt.empty and "FACILITY_ID" in trt.columns:
                keep_t = [c for c in ["FACILITY_ID","TREATMENT_OBJECTIVE_CODE","TREATMENT_PROCESS_CODE"] if c in trt.columns]
                t_min = trt[keep_t].drop_duplicates()
                tp_df = tp_df.merge(t_min, on="FACILITY_ID", how="left")
            # sort by facility_name
            if "FACILITY_NAME" in tp_df.columns:
                tp_df = tp_df.sort_values(["FACILITY_NAME"], kind="mergesort")
            for _, r in tp_df.iterrows():
                tr_rows.append([
                    r.get("FACILITY_NAME", ""),
                    active_from(r.get("FACILITY_ACTIVITY_CODE", "")),
                    r.get("FACILITY_ID", ""),
                    r.get("STATE_FACILITY_ID", ""),
                    desc("TREATMENT_OBJECTIVE_CODE", r.get("TREATMENT_OBJECTIVE_CODE", "")),
                    r.get("TREATMENT_PROCESS_CODE", ""),
                ])
    if tr_rows:
        add_table(doc,
                  headers=["Name", "Active?", "SDWIS Facility ID", "State Facility ID", "Treatment Objective", "Treatment Process"],
                  rows=tr_rows)
    else:
        doc.add_paragraph("No data available.")

# -------- Storage (codes + name fallback; exclude sources; sort by facility_name)
    stor_rows = []
    if not wsf.empty:
        df = wsf.copy()
    
        # normalize fields we rely on
        for c in ("FACILITY_TYPE_CODE","FACILITY_NAME","FACILITY_ACTIVITY_CODE","IS_SOURCE_IND"):
            if c in df.columns:
                df[c] = df[c].astype(str)
                
        print("FACILITY_TYPE_CODE unique:",
              sorted(df["FACILITY_TYPE_CODE"].dropna().str.upper().unique().tolist()))
    
        # exclude obvious sources
        not_source = ~df.get("IS_SOURCE_IND", "").str.upper().eq("Y")
    
        # primary: match SDWIS storage-ish facility types
        storage_codes = {"ST", "CW", "RS"}  # Storage, Clear Well, Reservoir
        code_hit = df.get("FACILITY_TYPE_CODE", "").str.upper().isin(storage_codes)
    
        # fallback: common storage keywords in the name
        # (kept simple; you can expand this list if you see more patterns)
        name_pattern = r"\b(TANK|TANKS|CLEARWELL|CLEAR WELL|RESERVOIR|STANDPIPE|ELEVATED|GROUND STORAGE|GST|EST|CST|TOWER)\b"
        name_hit = df.get("FACILITY_NAME", "").str.contains(name_pattern, case=False, na=False)
    
        sd = df[not_source & (code_hit | name_hit)].copy()
    
        if not sd.empty and "FACILITY_NAME" in sd.columns:
            sd = sd.sort_values(["FACILITY_NAME"], kind="mergesort")
    
        for _, r in sd.iterrows():
            stor_rows.append([
                r.get("FACILITY_NAME", ""),
                active_from(r.get("FACILITY_ACTIVITY_CODE", "")),
                r.get("FACILITY_ID", ""),
                r.get("STATE_FACILITY_ID", ""),
            ])
    
    if stor_rows:
        doc.add_paragraph("")  # spacer
        doc.add_heading("Storage", level=2)
        add_table(doc,
                  headers=["Name", "Active?", "SDWIS Facility ID", "State Facility ID"],
                  rows=stor_rows)
    # else: omit Storage when none found

    # ======================== Violations ========================
    doc.add_heading("Violations", level=1)

    def vio_rows_from(df: pd.DataFrame) -> list[list[str]]:
        rows = []
        for _, r in df.iterrows():
            rows.append([
                desc("VIOLATION_CATEGORY_CODE", r.get("VIOLATION_CATEGORY_CODE", "")),
                desc("VIOLATION_CODE", r.get("VIOLATION_CODE", "")),
                r.get("CONTAMINANT_CODE", ""),
            ])
        return rows

    # Health Based: only Y; sort by category then code
    doc.add_paragraph("Health Based")
    hb_rows = []
    if not vio.empty and "IS_HEALTH_BASED_IND" in vio.columns:
        hb = vio[vio["IS_HEALTH_BASED_IND"].astype(str).str.upper().eq("Y")].copy()
        sort_cols = [c for c in ["VIOLATION_CATEGORY_CODE","VIOLATION_CODE"] if c in hb.columns]
        if sort_cols:
            hb = hb.sort_values(sort_cols, kind="mergesort")
        hb_rows = vio_rows_from(hb)
    if hb_rows:
        add_table(doc, headers=["Category", "Type", "Contaminant"], rows=hb_rows)
    else:
        doc.add_paragraph("No data available.")

    # Non-Health Based: only N; sort by category then code
    doc.add_paragraph("")  # spacer
    doc.add_paragraph("Non-Health Based")
    nh_rows = []
    if not vio.empty and "IS_HEALTH_BASED_IND" in vio.columns:
        nh = vio[vio["IS_HEALTH_BASED_IND"].astype(str).str.upper().eq("N")].copy()
        sort_cols = [c for c in ["VIOLATION_CATEGORY_CODE","VIOLATION_CODE"] if c in nh.columns]
        if sort_cols:
            nh = nh.sort_values(sort_cols, kind="mergesort")
        nh_rows = vio_rows_from(nh)
    if nh_rows:
        add_table(doc, headers=["Category", "Type", "Contaminant"], rows=nh_rows)
    else:
        doc.add_paragraph("No data available.")

    # ---------------- save ----------------
    if out_path is None:
        out_path = f"{pwsid}_SDWIS_Report.docx"
    doc.save(out_path)
    print(f"Report saved: {out_path}")
    return out_path

import os
import tempfile
import zipfile

def generate_reports_zip(pwsids: list[str], fetch_data_fn, zip_name: str = "SDWIS_Reports.zip") -> str:
    """
    Create multiple Word reports and bundle them into a ZIP.

    Parameters
    ----------
    pwsids : list[str]
        PWSIDs to include.
    fetch_data_fn : callable
        Function that accepts (pwsid) -> dict[str, pd.DataFrame]
        e.g., your cached_fetch_all_selected from app.py.
    zip_name : str
        Filename for the resulting zip (only the name; we'll place it in a temp dir).

    Returns
    -------
    str
        Absolute path to the created ZIP file.
    """
    if not HAVE_DOCX:
        raise RuntimeError("python-docx is not installed. Install it with: pip install python-docx")

    tmpdir = tempfile.mkdtemp()
    report_paths: list[str] = []

    # Build each DOCX
    for pid in pwsids:
        try:
            data = fetch_data_fn(pid)
            out_docx = os.path.join(tmpdir, f"{pid}_SDWIS_Report.docx")
            generate_report(pid, data, out_path=out_docx)
            report_paths.append(out_docx)
        except Exception as e:
            print(f"❌ Failed to generate report for {pid}: {e}")

    if not report_paths:
        raise RuntimeError("No reports generated; ZIP will not be created.")

    # Bundle into ZIP (in the same temp dir)
    zip_path = os.path.join(tmpdir, zip_name)
    with zipfile.ZipFile(zip_path, "w", zipfile.ZIP_DEFLATED) as zf:
        for path in report_paths:
            zf.write(path, arcname=os.path.basename(path))

    return zip_path

# ----------------------- Main (state-agnostic CLI) -----------------------

def main():
    print("SDWIS Finder (state-aware)")
    print("  • Enter a PWSID (e.g., CA1010016), OR")
    print("  • Enter a 2-letter STATE code (e.g., CA), then (optional) name tokens and/or county")
    query = input("\nSearch (PWSID or STATE code): ").strip()

    if looks_like_pwsid(query):
        pwsid = query.upper()
    else:
        sc = query.strip().upper()
        if not re.fullmatch(r"[A-Z]{2}", sc):
            print("Please enter a valid PWSID or a 2-letter state code (e.g., CA, TX, NY).")
            sys.exit(0)

        name = input("Water system name (optional): ").strip()
        county_hint = input("County or City (optional): ").strip() or None
        print(f"\nSearching {sc}…")
        matches = search_by_name(sc, name_query=name, county_filter=county_hint)
        if matches.empty:
            print("No matches found. Try different words or adjust county filter.")
            sys.exit(0)

        # Show matches
        display = matches.reset_index(drop=True)
        display = display.reset_index().rename(columns={"index": "#"})
        print("\nMatches:")
        print(display.to_string(index=False))

        # Pick one
        try:
            pick = int(input("\nEnter the # of the system to use: ").strip())
            pwsid = matches.iloc[pick]["PWSID"]
        except Exception:
            print("Invalid selection.")
            sys.exit(1)

    print(f"\nFetching selected fields for PWSID: {pwsid}")
    all_data = fetch_all_selected(pwsid)

    if not HAVE_DOCX:
        print("Skipping Word report because python-docx is not installed.")
        sys.exit(0)

    generate_report(pwsid, all_data)

if __name__ == "__main__":
    main()
